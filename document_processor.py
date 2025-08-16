from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import tempfile
import os
import base64
from typing import List, Optional
from langchain.schema import Document
from docx import Document as DocxDocument
from PIL import Image

class DocumentProcessor:
    def __init__(self):
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200
        )
    
    def load_document(self, file_path: str) -> List[Document]:
        """Load document from file path"""
        if file_path.endswith('.pdf'):
            loader = PyPDFLoader(file_path)
        elif file_path.endswith('.docx') or file_path.endswith('.doc'):
            return self._load_docx_document(file_path)
        elif self._is_image_file(file_path):
            return self._load_image_document(file_path)
        else:
            loader = TextLoader(file_path, encoding='utf-8')
        
        documents = loader.load()
        return self.text_splitter.split_documents(documents)
    
    def _load_docx_document(self, file_path: str) -> List[Document]:
        """Load Word document and extract text"""
        doc = DocxDocument(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        # Create a Document object
        document = Document(page_content=text, metadata={"source": file_path})
        return self.text_splitter.split_documents([document])
    
    def _is_image_file(self, file_path: str) -> bool:
        """Check if file is an image"""
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.tiff']
        return any(file_path.lower().endswith(ext) for ext in image_extensions)
    
    def _load_image_document(self, file_path: str) -> List[Document]:
        """Load image and encode as base64 for vision model processing"""
        with open(file_path, 'rb') as image_file:
            image_data = base64.b64encode(image_file.read()).decode('utf-8')
        
        # Create a Document object with image data
        document = Document(
            page_content=f"[IMAGE_DATA]{image_data}[/IMAGE_DATA]",
            metadata={"source": file_path, "type": "image"}
        )
        return [document]  # Don't split images
    
    def load_from_upload(self, uploaded_file) -> List[Document]:
        """Load document from uploaded file"""
        with tempfile.NamedTemporaryFile(
            delete=False, 
            suffix=f".{uploaded_file.name.split('.')[-1]}"
        ) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            chunks = self.load_document(tmp_file_path)
            return chunks
        finally:
            os.unlink(tmp_file_path)
    
    def create_vectorstore(self, chunks: List[Document]) -> FAISS:
        """Create FAISS vector store from chunks"""
        return FAISS.from_documents(chunks, self.embeddings)
    
    def search_documents(self, vectorstore: FAISS, query: str, k: int = 3) -> List[Document]:
        """Search for relevant documents"""
        return vectorstore.similarity_search(query, k=k)